# -*- coding: utf-8 -*-
from __future__ import absolute_import, unicode_literals

import json
import os

from celery import shared_task
from django.conf import settings
from django.shortcuts import get_object_or_404
from django.utils import timezone
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docxcompose.composer import Composer

from finance.models import JobContract
from finance.utils import word_to_pdf

from gearfarm.utils.common_utils import format_currency
from oauth.e_sign import ESign


@shared_task
def create_job_contract_docx_pdf_files(contract_id):
    """
    生成doc   预览pdf   下载pdf
    """
    generate_preview_contract(contract_id)
    generate_contract(contract_id)


@shared_task
def create_confidentiality_agreement_docx_pdf_files(contract_id):
    """
    生成doc   预览pdf   下载pdf
    """
    generate_preview_confidentiality_agreement(contract_id)
    generate_confidentiality_agreement(contract_id)


@shared_task
def generate_preview_contract(contract_id):
    job_contract = JobContract.objects.filter(pk=contract_id).first()
    if not job_contract:
        return

    current_dir = os.path.dirname(os.path.abspath(__file__))
    contract_path = os.path.join(current_dir, 'contract_template/job_contract_template.docx')
    contract_type = job_contract.contract_type
    contract_category = job_contract.contract_category
    if contract_type == 'design':
        contract_path = os.path.join(current_dir, 'contract_template/design_contract_template.docx')
    elif contract_type == 'regular':
        contract_path = os.path.join(current_dir, 'contract_template/regular_job_contract_template.docx')

    amount_in_words = format_currency(job_contract.contract_money)
    project_name = ''
    # 打款方式
    if contract_category == 'project':
        project = job_contract.project
        # 文本替换
        project_name = project.name
        if job_contract.pay_way == 'installments':
            remit_way = json.loads(job_contract.remit_way)
            remit_way_str = ''
            for index, i in enumerate(remit_way):
                remit_way_str += '({})第{}阶段：【{}】,甲方向乙方支付合同金额的【{}】%,共计【{}】元,即（大写）【{}】;\n'.format(index + 1, index + 1,
                                                                                                i['name'],
                                                                                                i['proportion'],
                                                                                                i['money'],
                                                                                                format_currency(
                                                                                                    i['money']))
            if contract_type == 'design':
                remit_way_str = '在乙方完成相应阶段工作成果后且严格遵守本协议约定的乙方各类义务的前提下，甲方将分阶段向乙方支付以下款项：\n' + remit_way_str
        else:
            remit_way_str = '合同总金额为人民币（大写）{}（￥{}元），为乙方经营所得款项。验收通过后，一次性支付此款项。'.format(
                amount_in_words, job_contract.contract_money)
    else:
        if job_contract.pay_way == 'installments':
            remit_way = json.loads(job_contract.remit_way)
            # {"timedelta_weeks": 2, "c": 9000}
            weekly_amount_in_words = format_currency(remit_way['amount'])
            remit_way_str = "项目进度款：乙方向甲方提交符合约定的工期的进度报告后，甲方将按【周】向乙方支付合同进度款，每【{}周】金额为人民币（大写）{}（￥{}元）。".format(
                remit_way['timedelta_weeks'], weekly_amount_in_words, remit_way['amount']
            )
        else:
            remit_way_str = '乙方向甲方提交符合约定的工期进度报告后，甲方向乙方一次性支付合同款，即人民币（大写）{}（￥{}元）。'.format(
                amount_in_words, job_contract.contract_money)

    today = timezone.now().date()
    sign_date_str = '{year}年{month}月{day}日'.format(year=today.year, month=today.month, day=today.day)

    manager = job_contract.manager
    manager_name = manager.username if manager else ''
    manager_email = manager.email if manager else ''
    manager_phone = manager.profile.phone if manager else ''

    doc = Document(contract_path)
    project_results_show = json.loads(job_contract.project_results_show)
    show_data = ','.join(i for i in project_results_show)
    fields_dict = {
        'contract_name': job_contract.contract_name,
        'developer_name': job_contract.name,
        'id_card_number': job_contract.id_card_number,
        'project_name': project_name,
        "develop_sprint": job_contract.develop_sprint,
        "amount_in_words": amount_in_words,
        'contract_money': job_contract.contract_money,
        'remit_way': remit_way_str,
        'payee_account': job_contract.payee_account,
        'payee_opening_bank': job_contract.payee_opening_bank,
        'payee_name': job_contract.payee_name,
        "maintain_period": job_contract.maintain_period,
        "manager_name": manager_name,
        "manager_email": manager_email,
        "manager_phone": manager_phone,
        'developer_email': job_contract.email,
        'developer_phone': job_contract.phone,
        'develop_date_start': job_contract.develop_date_start,
        'develop_date_end': job_contract.develop_date_end,
        'project_results_show': show_data,
        'style_confirm': job_contract.style_confirm,
        'global_design': job_contract.global_design,
        'walk_through': job_contract.walk_through,
        '签署日期': sign_date_str
    }
    for fields, values in fields_dict.items():
        for p in doc.paragraphs:
            if fields in p.text:
                inline = p.runs
                for i in inline:
                    if fields in i.text:
                        text = i.text.replace(fields, str(values))
                        i.text = text
    # 身份证文件、项目功能文档
    front_side_of_id_card = job_contract.front_side_of_id_card.file if job_contract.front_side_of_id_card else None
    back_side_of_id_card = job_contract.back_side_of_id_card.file if job_contract.back_side_of_id_card else None
    requirement_word = None
    if job_contract.develop_function_declaration:
        requirement_word = Document(job_contract.develop_function_declaration.file)
    delivery_list = None
    if job_contract.delivery_list:
        delivery_list = Document(job_contract.delivery_list.file)
    if front_side_of_id_card or back_side_of_id_card:
        for p in doc.paragraphs:
            if '工程师身份证扫描件' in p.text:
                if front_side_of_id_card:
                    p.runs[-1].add_picture(front_side_of_id_card, width=Inches(5))
                if back_side_of_id_card:
                    p.runs[-1].add_picture(back_side_of_id_card, width=Inches(5))
                break
    master = doc
    composer = Composer(master)
    if requirement_word:
        composer.append(requirement_word)
    if delivery_list:
        composer.append(delivery_list)
        # for requirement_p in requirement_word.paragraphs:
        #     for r in requirement_p.runs:
        #         r.font.name = 'Times New Roman'  # 控制是西文时的字体
        #         r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        #
        # for p_index, p in enumerate(doc._element.body):
        #     if hasattr(p, 'xml'):
        #         if '附件2: 项目功能需求模块说明' in p.xml:
        #             current_index = p_index + 1
        #             for sub_body in requirement_word._element.body:
        #                 doc._element.body.insert(current_index, sub_body)
        #                 current_index = current_index + 1
        #             break
        # for r in develop_function_declaration_word.paragraphs:
        #     p.runs[-1].add_break()
        #     run = p.add_run(r.text)
        #     font = run.font
        #     font.name = '宋体'
        #     font.size = Pt(14)
    file_path = settings.MEDIA_ROOT + 'finance/developer_contract/{}_contract_preview.docx'.format(job_contract.id)
    composer.save(file_path)
    dir_path = os.path.dirname(file_path)
    if not os.path.exists(dir_path):
        os.makedirs(dir_path)
    # doc.save(file_path)
    word_to_pdf(file_path, contract_id)


@shared_task
def generate_contract(contract_id):
    job_contract = get_object_or_404(JobContract, pk=contract_id)
    in_file_path = settings.MEDIA_ROOT + 'finance/developer_contract/{}_contract_preview.docx'.format(job_contract.id)
    doc = Document(in_file_path)
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0, 0, 0)
    out_file_path = settings.MEDIA_ROOT + 'finance/developer_contract/{}_contract.docx'.format(job_contract.id)
    doc.save(out_file_path)
    word_to_pdf(out_file_path, contract_id)


def date_chinese_strftime(date):
    return date.strftime("%Y{y}%m{m}%d{d}").format(y="年", m="月", d="日")


@shared_task
def generate_preview_confidentiality_agreement(contract_id):
    today = timezone.now().date()
    sign_date_str = '{year}年{month}月{day}日'.format(year=today.year, month=today.month, day=today.day)
    job_contract = get_object_or_404(JobContract, pk=contract_id)
    current_dir = os.path.dirname(os.path.abspath(__file__))
    contract_path = os.path.join(current_dir, 'contract_template/job_confidentiality_agreement_template.docx')
    manager = job_contract.manager
    manager_name = manager.username if manager else ''
    manager_email = manager.email if manager else ''
    doc = Document(contract_path)
    fields_dict = {
        '项目经理姓名': manager_name,
        '项目经理邮箱': manager_email,

        '工程师姓名': job_contract.name,
        '工程师联系地址': job_contract.address,
        'E': job_contract.email,
        '提交时间': date_chinese_strftime(
            job_contract.committed_at) if job_contract.committed_at else date_chinese_strftime(job_contract.created_at),
        'C': job_contract.contract_name,
        '签署日期': sign_date_str
    }
    # 替换docx内容
    for fields, values in fields_dict.items():
        for p in doc.paragraphs:
            if fields in p.text:
                inline = p.runs
                for i in inline:
                    if fields in i.text:
                        text = i.text.replace(fields, str(values))
                        i.text = text

    file_path = settings.MEDIA_ROOT + 'finance/developer_contract/{}_confidentiality_agreement_preview.docx'.format(
        job_contract.id)
    dir_path = os.path.dirname(file_path)
    if not os.path.exists(dir_path):
        os.makedirs(dir_path)
    doc.save(file_path)
    word_to_pdf(file_path, contract_id)


@shared_task
def generate_confidentiality_agreement(contract_id):
    job_contract = get_object_or_404(JobContract, pk=contract_id)
    in_file_path = settings.MEDIA_ROOT + 'finance/developer_contract/{}_confidentiality_agreement_preview.docx'.format(
        job_contract.id)
    doc = Document(in_file_path)
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0, 0, 0)
    out_file_path = settings.MEDIA_ROOT + 'finance/developer_contract/{}_confidentiality_agreement.docx'.format(
        job_contract.id)
    doc.save(out_file_path)
    word_to_pdf(out_file_path, contract_id)


@shared_task
def query_sign_result():
    job_contracts = JobContract.objects.filter(is_esign_contract=True, status='waiting')
    for job_contract in job_contracts:
        flow_id = job_contract.flow_id
        secret_flow_id = job_contract.secret_flow_id
        code, msg, data = ESign.sign_flow_query(flow_id)
        if data['flowStatus'] == 2:
            job_contract.is_sign_contract = True
            job_contract.signed_at = timezone.now()
        code1, msg1, data1 = ESign.sign_flow_query(secret_flow_id)
        if data1['flowStatus'] == 2:
            job_contract.is_sign_secret = True
        if job_contract.is_sign_secret and job_contract.is_sign_contract:
            job_contract.status = 'signed'

        job_contract.save()
